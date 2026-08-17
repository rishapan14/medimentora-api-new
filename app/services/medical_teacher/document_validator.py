"""Validation for AI Medical Teacher document uploads (PDF / DOCX / TXT)."""

from __future__ import annotations

import hashlib
import logging
from dataclasses import dataclass, field
from typing import Iterable

from werkzeug.datastructures import FileStorage

logger = logging.getLogger(__name__)

ALLOWED_EXTENSIONS = frozenset({"pdf", "docx", "txt"})
ALLOWED_MIME_TYPES = frozenset(
  {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
    "text/plain",
    "application/octet-stream",
  }
)

PDF_MAGIC = b"%PDF"
ZIP_MAGIC = b"PK"  # DOCX is a ZIP package
DOCX_CONTENT_TYPES = b"[Content_Types].xml"


@dataclass
class DocumentValidationIssue:
  filename: str
  code: str
  message: str


@dataclass
class ValidatedDocument:
  storage: FileStorage
  filename: str
  extension: str
  file_type: str  # pdf | docx | txt
  size_bytes: int
  content_hash: str
  mime_type: str | None


@dataclass
class DocumentValidationResult:
  ok: bool
  files: list[ValidatedDocument] = field(default_factory=list)
  errors: list[DocumentValidationIssue] = field(default_factory=list)
  total_size_bytes: int = 0

  def error_messages(self) -> list[str]:
    return [f"{e.filename}: {e.message}" for e in self.errors]


class TeacherDocumentValidator:
  """Validate medical textbook / notes uploads before persistence."""

  def __init__(
    self,
    max_files: int = 5,
    max_total_bytes: int = 200 * 1024 * 1024,
    max_file_bytes: int | None = None,
  ):
    self.max_files = max_files
    self.max_total_bytes = max_total_bytes
    self.max_file_bytes = max_file_bytes or max_total_bytes

  def validate(self, files: Iterable[FileStorage | None]) -> DocumentValidationResult:
    result = DocumentValidationResult(ok=True)
    seen_hashes: set[str] = set()
    seen_names: set[str] = set()

    raw_files = [f for f in files if f is not None]
    if not raw_files:
      result.ok = False
      result.errors.append(
        DocumentValidationIssue(
          filename="(none)",
          code="no_files",
          message="At least one PDF, DOCX, or TXT file is required.",
        )
      )
      return result

    if len(raw_files) > self.max_files:
      result.ok = False
      result.errors.append(
        DocumentValidationIssue(
          filename="(batch)",
          code="too_many_files",
          message=f"Maximum {self.max_files} files allowed per upload (received {len(raw_files)}).",
        )
      )
      return result

    for storage in raw_files:
      item = self._validate_one(storage, seen_hashes, seen_names)
      if isinstance(item, DocumentValidationIssue):
        result.errors.append(item)
        continue
      result.files.append(item)
      result.total_size_bytes += item.size_bytes

    if result.total_size_bytes > self.max_total_bytes:
      result.ok = False
      mb = self.max_total_bytes / (1024 * 1024)
      result.errors.append(
        DocumentValidationIssue(
          filename="(batch)",
          code="total_too_large",
          message=f"Total upload size exceeds {mb:.0f} MB.",
        )
      )

    if result.errors:
      result.ok = False
    return result

  def _validate_one(
    self,
    storage: FileStorage,
    seen_hashes: set[str],
    seen_names: set[str],
  ) -> ValidatedDocument | DocumentValidationIssue:
    filename = (storage.filename or "").strip()
    if not filename:
      return DocumentValidationIssue("", "missing_filename", "Filename is missing.")

    name_key = filename.lower()
    if name_key in seen_names:
      return DocumentValidationIssue(filename, "duplicate_name", "Duplicate filename in this upload.")
    seen_names.add(name_key)

    extension = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if extension not in ALLOWED_EXTENSIONS:
      return DocumentValidationIssue(
        filename,
        "unsupported_type",
        f"Unsupported file type '.{extension}'. Allowed: PDF, DOCX, TXT.",
      )

    mime = (storage.mimetype or "").lower() or None
    if mime and mime not in ALLOWED_MIME_TYPES:
      logger.warning("Unusual MIME type %s for %s — continuing with extension check", mime, filename)

    raw = storage.stream.read()
    storage.stream.seek(0)
    size = len(raw)
    if size == 0:
      return DocumentValidationIssue(filename, "empty_file", "File is empty.")
    if size > self.max_file_bytes:
      mb = self.max_file_bytes / (1024 * 1024)
      return DocumentValidationIssue(
        filename,
        "file_too_large",
        f"File exceeds maximum size of {mb:.0f} MB.",
      )

    content_hash = hashlib.sha256(raw).hexdigest()
    if content_hash in seen_hashes:
      return DocumentValidationIssue(filename, "duplicate_content", "Duplicate file content in this upload.")
    seen_hashes.add(content_hash)

    integrity = self._check_integrity(filename, extension, raw)
    if integrity:
      return integrity

    file_type = extension if extension != "doc" else "docx"
    return ValidatedDocument(
      storage=storage,
      filename=filename,
      extension=extension,
      file_type=file_type,
      size_bytes=size,
      content_hash=content_hash,
      mime_type=mime,
    )

  def _check_integrity(
    self,
    filename: str,
    extension: str,
    raw: bytes,
  ) -> DocumentValidationIssue | None:
    if extension == "pdf":
      if not raw.startswith(PDF_MAGIC):
        return DocumentValidationIssue(filename, "invalid_pdf", "File is not a valid PDF.")
      try:
        import fitz

        doc = fitz.open(stream=raw, filetype="pdf")
        try:
          if doc.is_encrypted:
            return DocumentValidationIssue(filename, "encrypted_pdf", "Encrypted PDFs are not supported.")
          if doc.page_count < 1:
            return DocumentValidationIssue(filename, "empty_pdf", "PDF has no pages.")
        finally:
          doc.close()
      except Exception as exc:
        logger.warning("PDF validation failed for %s: %s", filename, exc)
        return DocumentValidationIssue(filename, "corrupt_pdf", "PDF could not be opened (corrupt or unreadable).")

    elif extension == "docx":
      if not raw.startswith(ZIP_MAGIC):
        return DocumentValidationIssue(filename, "invalid_docx", "File is not a valid DOCX (ZIP) package.")
      if DOCX_CONTENT_TYPES not in raw[:200_000] and DOCX_CONTENT_TYPES not in raw:
        # Soft check — many DOCX include Content_Types early
        try:
          from io import BytesIO

          from docx import Document

          Document(BytesIO(raw))
        except ImportError:
          return DocumentValidationIssue(
            filename,
            "docx_dependency_missing",
            "DOCX support requires python-docx. Install it on the API server.",
          )
        except Exception:
          return DocumentValidationIssue(filename, "corrupt_docx", "DOCX could not be opened.")
      else:
        try:
          from io import BytesIO

          from docx import Document

          Document(BytesIO(raw))
        except ImportError:
          return DocumentValidationIssue(
            filename,
            "docx_dependency_missing",
            "DOCX support requires python-docx. Install it on the API server.",
          )
        except Exception:
          return DocumentValidationIssue(filename, "corrupt_docx", "DOCX could not be opened.")

    elif extension == "txt":
      try:
        raw.decode("utf-8")
      except UnicodeDecodeError:
        try:
          raw.decode("latin-1")
        except Exception:
          return DocumentValidationIssue(filename, "invalid_txt", "TXT file encoding is not readable.")

    return None
