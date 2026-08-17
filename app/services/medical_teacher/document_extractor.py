"""Document text extraction for AI Medical Teacher (Module 1).

Supports:
- PDF: native text via PyMuPDF; OCR fallback for scanned pages (reuses OCRService)
- DOCX: python-docx paragraphs + tables
- TXT: UTF-8 / latin-1

Preserves lightweight structure (pages, headings, lists, page numbers).
Full chapter/topic split belongs to Module 2 (Book Parser).
"""

from __future__ import annotations

import logging
import math
import os
import re
from dataclasses import dataclass, field
from pathlib import Path

logger = logging.getLogger(__name__)

MIN_DIGITAL_CHARS_PER_PAGE = 40
HEADING_MAX_LEN = 120


@dataclass
class PageStructure:
  page_number: int
  text: str
  headings: list[str] = field(default_factory=list)
  lists: list[str] = field(default_factory=list)
  tables_count: int = 0
  images_count: int = 0
  is_scanned: bool = False
  engine: str | None = None
  confidence: float | None = None

  def to_dict(self) -> dict:
    return {
      "page_number": self.page_number,
      "text": self.text,
      "headings": self.headings,
      "lists": self.lists[:50],
      "tables_count": self.tables_count,
      "images_count": self.images_count,
      "is_scanned": self.is_scanned,
      "engine": self.engine,
      "confidence": self.confidence,
      "char_count": len(self.text or ""),
    }


@dataclass
class DocumentExtractionResult:
  success: bool
  text: str = ""
  method: str | None = None
  page_count: int = 0
  char_count: int = 0
  word_count: int = 0
  confidence: float | None = None
  pages: list[PageStructure] = field(default_factory=list)
  message: str = ""
  error_code: str | None = None

  def structure_dict(self) -> dict:
    return {
      "page_count": self.page_count,
      "extraction_method": self.method,
      "pages": [p.to_dict() for p in self.pages],
      "heading_count": sum(len(p.headings) for p in self.pages),
      "note": "Lightweight structure from Document Processing (Module 1). Chapter split is Module 2.",
    }

  def to_dict(self) -> dict:
    return {
      "success": self.success,
      "method": self.method,
      "page_count": self.page_count,
      "char_count": self.char_count,
      "word_count": self.word_count,
      "confidence": self.confidence,
      "message": self.message,
      "error_code": self.error_code,
      "text_preview": (self.text or "")[:500],
      "structure": self.structure_dict() if self.success else None,
    }


class DocumentExtractor:
  """Extract educational text and structure from uploaded medical documents."""

  @classmethod
  def extract(cls, file_path: str, file_type: str, progress_callback=None) -> DocumentExtractionResult:
    path = Path(file_path)
    if not path.is_file():
      return DocumentExtractionResult(
        success=False,
        message="Document file not found on disk.",
        error_code="missing_file",
      )

    normalized = (file_type or "").strip().lower()
    try:
      if normalized == "pdf":
        return cls._extract_pdf(str(path), progress_callback=progress_callback)
      if normalized == "docx":
        cls._report_progress(progress_callback, "extracting_content", 35)
        result = cls._extract_docx(str(path))
        cls._report_progress(progress_callback, "cleaning_content", 85)
        return result
      if normalized == "txt":
        cls._report_progress(progress_callback, "extracting_content", 50)
        result = cls._extract_txt(str(path))
        cls._report_progress(progress_callback, "cleaning_content", 85)
        return result
      return DocumentExtractionResult(
        success=False,
        message=f"Unsupported file type for extraction: {file_type}",
        error_code="unsupported_type",
      )
    except Exception:
      logger.exception("Document extraction failed for %s", file_path)
      return DocumentExtractionResult(
        success=False,
        message="Document extraction failed unexpectedly.",
        error_code="extraction_failed",
      )

  # --- PDF -----------------------------------------------------------------

  @classmethod
  def _extract_pdf(cls, file_path: str, progress_callback=None) -> DocumentExtractionResult:
    try:
      import fitz
    except ImportError:
      return DocumentExtractionResult(
        success=False,
        message="PyMuPDF is required for PDF extraction.",
        error_code="pymupdf_missing",
      )

    doc = fitz.open(file_path)
    pages: list[PageStructure] = []
    confidences: list[float] = []
    any_ocr = False
    any_native = False

    try:
      total_pages = max(1, len(doc))
      cls._report_progress(progress_callback, "extracting_content", 20)
      for idx in range(total_pages):
        page_num = idx + 1
        page = doc.load_page(idx)
        headings, lists, tables_count, images_count = cls._pdf_page_structure(page)
        digital = (page.get_text("text") or "").strip()

        if len(digital) >= MIN_DIGITAL_CHARS_PER_PAGE:
          any_native = True
          if not headings:
            headings = cls._detect_headings_from_text(digital)
          if not lists:
            lists = cls._detect_lists_from_text(digital)
          pages.append(
            PageStructure(
              page_number=page_num,
              text=digital,
              headings=headings,
              lists=lists,
              tables_count=tables_count,
              images_count=images_count,
              is_scanned=False,
              engine="pymupdf",
              confidence=0.95,
            )
          )
          confidences.append(0.95)
          cls._report_progress(
            progress_callback,
            "extracting_content",
            20 + round(((idx + 1) / total_pages) * 60),
          )
          continue

        # Scanned / low-text page → OCR via existing medical OCR pipeline
        any_ocr = True
        ocr_text, confidence, engine = cls._ocr_pdf_page(page)
        if not headings:
          headings = cls._detect_headings_from_text(ocr_text)
        if not lists:
          lists = cls._detect_lists_from_text(ocr_text)
        pages.append(
          PageStructure(
            page_number=page_num,
            text=ocr_text or "",
            headings=headings,
            lists=lists,
            tables_count=tables_count,
            images_count=images_count,
            is_scanned=True,
            engine=engine,
            confidence=confidence,
          )
        )
        if confidence is not None:
          confidences.append(confidence)
        cls._report_progress(
          progress_callback,
          "ocr_processing",
          20 + round(((idx + 1) / total_pages) * 60),
        )
    finally:
      doc.close()

    cls._report_progress(progress_callback, "cleaning_content", 85)
    cls._clean_repeated_page_noise(pages)
    merged = cls._merge_pages(pages)
    if not merged.strip():
      return DocumentExtractionResult(
        success=False,
        message="No readable text found in this PDF (may be image-only without OCR engines).",
        error_code="no_text",
        page_count=len(pages),
        pages=pages,
      )

    method = "hybrid" if (any_ocr and any_native) else ("ocr" if any_ocr else "native_pdf")
    avg_conf = sum(confidences) / len(confidences) if confidences else None
    return cls._finalize(merged, method, pages, avg_conf)

  @classmethod
  def _pdf_page_structure(cls, page) -> tuple[list[str], list[str], int, int]:
    """Derive headings/lists/table/image counts from PyMuPDF dict blocks."""
    headings: list[str] = []
    lists: list[str] = []
    tables_count = 0
    images_count = 0

    try:
      images_count = len(page.get_images(full=True) or [])
    except Exception:
      images_count = 0

    try:
      # PyMuPDF table finder (available in recent versions)
      finder = getattr(page, "find_tables", None)
      if callable(finder):
        tables = finder()
        tables_count = len(getattr(tables, "tables", []) or [])
    except Exception:
      tables_count = 0

    try:
      data = page.get_text("dict") or {}
      font_sizes: list[float] = []
      spans_meta: list[tuple[float, str, bool]] = []
      for block in data.get("blocks", []):
        if block.get("type") != 0:
          continue
        for line in block.get("lines", []):
          parts = []
          max_size = 0.0
          is_bold = False
          for span in line.get("spans", []):
            text = (span.get("text") or "").strip()
            if not text:
              continue
            size = float(span.get("size") or 0)
            max_size = max(max_size, size)
            font_sizes.append(size)
            flags = int(span.get("flags") or 0)
            if flags & 2 ** 4:  # bold bit in PyMuPDF
              is_bold = True
            parts.append(text)
          line_text = " ".join(parts).strip()
          if line_text:
            spans_meta.append((max_size, line_text, is_bold))

      if font_sizes and spans_meta:
        avg = sum(font_sizes) / len(font_sizes)
        threshold = max(avg * 1.25, avg + 1.5)
        for size, text, is_bold in spans_meta:
          if len(text) <= HEADING_MAX_LEN and (size >= threshold or (is_bold and size >= avg)):
            if text not in headings:
              headings.append(text)

      for _, text, _ in spans_meta:
        if cls._looks_like_list_item(text) and text not in lists:
          lists.append(text)
    except Exception:
      logger.debug("Could not parse PDF dict structure for page", exc_info=True)

    return headings[:40], lists[:50], tables_count, images_count

  @classmethod
  def _ocr_pdf_page(cls, page) -> tuple[str, float | None, str]:
    """Render a PDF page and run the shared OCRService image pipeline."""
    import tempfile

    import fitz

    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
    tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
    tmp_path = tmp.name
    tmp.close()
    try:
      pix.save(tmp_path)
      from app.services.report_analysis.ocr.service import OCRService

      result = OCRService.extract_text(tmp_path, "image")
      if result.success and result.text:
        conf = None
        if result.quality and getattr(result.quality, "confidence", None) is not None:
          conf = float(result.quality.confidence)
        return result.text.strip(), conf, result.engine or "ocr"
      return "", None, result.engine or "ocr"
    finally:
      try:
        os.unlink(tmp_path)
      except OSError:
        pass

  # --- DOCX ----------------------------------------------------------------

  @classmethod
  def _extract_docx(cls, file_path: str) -> DocumentExtractionResult:
    try:
      from docx import Document
    except ImportError:
      return DocumentExtractionResult(
        success=False,
        message="DOCX support requires python-docx. Install: pip install python-docx",
        error_code="docx_dependency_missing",
      )

    document = Document(file_path)
    headings: list[str] = []
    lists: list[str] = []
    body_parts: list[str] = []

    for para in document.paragraphs:
      text = (para.text or "").strip()
      if not text:
        continue
      style_name = ""
      try:
        style_name = (para.style.name or "") if para.style else ""
      except Exception:
        style_name = ""

      if style_name.lower().startswith("heading") or style_name.lower().startswith("title"):
        headings.append(text)
        body_parts.append(f"\n## {text}\n")
      elif cls._looks_like_list_item(text) or style_name.lower().startswith("list"):
        lists.append(text)
        body_parts.append(f"• {text}")
      else:
        body_parts.append(text)

    tables_count = len(document.tables)
    for table in document.tables:
      rows = []
      for row in table.rows:
        cells = [((c.text or "").strip()) for c in row.cells]
        rows.append(" | ".join(cells))
      if rows:
        body_parts.append("\n[Table]\n" + "\n".join(rows) + "\n[/Table]\n")

    # Images count (approximate — relationships)
    images_count = 0
    try:
      images_count = len(document.inline_shapes)
    except Exception:
      images_count = 0

    merged = "\n".join(body_parts).strip()
    if not merged:
      return DocumentExtractionResult(
        success=False,
        message="No readable text found in this DOCX.",
        error_code="no_text",
      )

    page = PageStructure(
      page_number=1,
      text=merged,
      headings=headings,
      lists=lists,
      tables_count=tables_count,
      images_count=images_count,
      is_scanned=False,
      engine="python-docx",
      confidence=1.0,
    )
    return cls._finalize(merged, "docx", [page], 1.0)

  # --- TXT -----------------------------------------------------------------

  @classmethod
  def _extract_txt(cls, file_path: str) -> DocumentExtractionResult:
    raw = Path(file_path).read_bytes()
    text = None
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
      try:
        text = raw.decode(encoding)
        break
      except UnicodeDecodeError:
        continue
    if text is None:
      return DocumentExtractionResult(
        success=False,
        message="Could not decode TXT file.",
        error_code="invalid_txt",
      )

    text = text.lstrip("\ufeff").strip()
    headings = cls._detect_headings_from_text(text)
    lists = cls._detect_lists_from_text(text)
    page = PageStructure(
      page_number=1,
      text=text,
      headings=headings,
      lists=lists,
      is_scanned=False,
      engine="txt",
      confidence=1.0,
    )
    return cls._finalize(text, "txt", [page], 1.0)

  # --- Helpers -------------------------------------------------------------

  @classmethod
  def _finalize(
    cls,
    text: str,
    method: str,
    pages: list[PageStructure],
    confidence: float | None,
  ) -> DocumentExtractionResult:
    cleaned = cls._normalize_whitespace(text)
    words = len(re.findall(r"\b\w+\b", cleaned))
    return DocumentExtractionResult(
      success=True,
      text=cleaned,
      method=method,
      page_count=len(pages),
      char_count=len(cleaned),
      word_count=words,
      confidence=confidence,
      pages=pages,
      message="Document text extracted successfully.",
    )

  @staticmethod
  def _merge_pages(pages: list[PageStructure]) -> str:
    parts = []
    for page in pages:
      header = f"\n\n===== PAGE {page.page_number} =====\n\n"
      parts.append(header + (page.text or ""))
    return "".join(parts).strip()

  @staticmethod
  def _normalize_whitespace(text: str) -> str:
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"(?<=\w)-\n(?=[a-z])", "", text)
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{4,}", "\n\n\n", text)
    return text.strip()

  @classmethod
  def _clean_repeated_page_noise(cls, pages: list[PageStructure]) -> None:
    """Remove repeated headers, footers, page numbers, and adjacent OCR duplicates."""
    if not pages:
      return

    candidate_counts: dict[str, int] = {}
    page_candidates: list[set[str]] = []
    for page in pages:
      lines = [line.strip() for line in (page.text or "").splitlines() if line.strip()]
      keys = {
        cls._noise_key(line)
        for line in lines[:3] + lines[-3:]
        if 1 <= len(line) <= 160
        and cls._looks_like_running_noise(line)
        and cls._noise_key(line)
      }
      page_candidates.append(keys)
      for key in keys:
        candidate_counts[key] = candidate_counts.get(key, 0) + 1

    threshold = max(3, math.ceil(len(pages) * 0.5))
    repeated = {key for key, count in candidate_counts.items() if count >= threshold}

    for page, candidates in zip(pages, page_candidates):
      cleaned_lines: list[str] = []
      previous_key = ""
      for raw_line in (page.text or "").splitlines():
        line = raw_line.strip()
        key = cls._noise_key(line)
        is_page_number = bool(re.fullmatch(r"(?:page\s+)?\d+(?:\s+of\s+\d+)?", line, re.I))
        if key and key in repeated and key in candidates:
          continue
        if is_page_number or (key and key == previous_key):
          continue
        cleaned_lines.append(raw_line.rstrip())
        if key:
          previous_key = key
      page.text = cls._normalize_whitespace("\n".join(cleaned_lines))

  @staticmethod
  def _noise_key(value: str) -> str:
    value = re.sub(r"\s+", " ", (value or "").strip().lower())
    return value if any(char.isalpha() for char in value) else ""

  @staticmethod
  def _looks_like_running_noise(value: str) -> bool:
    line = (value or "").strip()
    lower = line.lower()
    if not line:
      return False
    if any(marker in lower for marker in ("copyright", "©", "confidential", "www.", "http")):
      return True
    if line.endswith((".", "?", "!", ":", ";")):
      return False
    return len(line.split()) <= 14

  @staticmethod
  def _report_progress(callback, stage: str, percent: int) -> None:
    if callback is None:
      return
    try:
      callback(stage, max(0, min(100, int(percent))))
    except Exception:
      logger.debug("Document progress callback failed", exc_info=True)

  @classmethod
  def _detect_headings_from_text(cls, text: str) -> list[str]:
    headings = []
    for line in (text or "").splitlines():
      line = line.strip()
      if not line or len(line) > HEADING_MAX_LEN:
        continue
      if re.match(r"^(chapter|unit|module|section|lesson)\s+\d+", line, re.I):
        headings.append(line)
      elif re.match(r"^\d+(\.\d+)*\s+[A-Z].{3,}", line) and len(line) < 80:
        headings.append(line)
      elif line.isupper() and 4 <= len(line) <= 80 and any(c.isalpha() for c in line):
        headings.append(line)
      elif line.startswith("#"):
        headings.append(line.lstrip("# ").strip())
    # de-dupe preserve order
    seen = set()
    out = []
    for h in headings:
      key = h.lower()
      if key not in seen:
        seen.add(key)
        out.append(h)
    return out[:40]

  @classmethod
  def _detect_lists_from_text(cls, text: str) -> list[str]:
    items = []
    for line in (text or "").splitlines():
      line = line.strip()
      if cls._looks_like_list_item(line):
        items.append(line)
    return items[:50]

  @staticmethod
  def _looks_like_list_item(text: str) -> bool:
    if not text or len(text) > 300:
      return False
    return bool(
      re.match(r"^([•\-–*]|\d+[\.)]|[a-zA-Z][\.)])\s+\S", text)
    )
